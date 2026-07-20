"""本地文件解析器 — 支持 PDF / DOCX / TXT / Excel (xlsx/xls)。
从用户上传的竞赛通知文件中提取原始文本，转为标准 raw_item 格式。
"""

import os
import json
import logging
from datetime import datetime
from typing import Optional

logger = logging.getLogger(__name__)

# 懒加载，避免未安装时阻塞整个模块
_pdf_reader = None
_docx_reader = None
_xlsx_reader = None


def _get_pdf_reader():
    global _pdf_reader
    if _pdf_reader is None:
        try:
            import PyPDF2
            _pdf_reader = PyPDF2
        except ImportError:
            try:
                import pypdf
                _pdf_reader = pypdf
            except ImportError:
                raise ImportError("PDF 解析需要安装 PyPDF2 或 pypdf: pip install PyPDF2")
    return _pdf_reader


def _get_docx_reader():
    global _docx_reader
    if _docx_reader is None:
        try:
            import docx
            _docx_reader = docx
        except ImportError:
            raise ImportError("Word 解析需要安装 python-docx: pip install python-docx")
    return _docx_reader


def _get_xlsx_reader():
    global _xlsx_reader
    if _xlsx_reader is None:
        try:
            import openpyxl
            _xlsx_reader = openpyxl
        except ImportError:
            raise ImportError("Excel 解析需要安装 openpyxl: pip install openpyxl")
    return _xlsx_reader


def _strip_name(path: str) -> str:
    """从文件路径中提取不带扩展名的文件名。"""
    name = os.path.basename(path)
    name = os.path.splitext(name)[0]
    return name


def _file_mtime(path: str) -> str:
    """文件的最后修改时间，ISO 格式。"""
    try:
        ts = os.path.getmtime(path)
        return datetime.fromtimestamp(ts).isoformat()
    except OSError:
        return ""


def parse_pdf(file_path: str) -> str:
    """解析 PDF，返回所有页面纯文本。"""
    reader = _get_pdf_reader()
    pages = []
    with open(file_path, "rb") as f:
        doc = reader.PdfReader(f)
        for page in doc.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
    return "\n\n".join(pages)


def parse_docx(file_path: str) -> str:
    """解析 Word 文档，返回纯文本。"""
    docx = _get_docx_reader()
    doc = docx.Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # 也读取表格
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n".join(paragraphs)


def parse_txt(file_path: str) -> str:
    """解析纯文本文件（支持 GBK/UTF-8 自动检测）。"""
    for enc in ["utf-8", "gbk", "gb2312", "latin-1"]:
        try:
            with open(file_path, "r", encoding=enc) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    raise ValueError(f"无法识别文件编码: {file_path}")


def parse_excel(file_path: str) -> str:
    """解析 Excel 所有 sheet，每个 sheet 转为表格文本。"""
    xl = _get_xlsx_reader()
    wb = xl.load_workbook(file_path, data_only=True)
    parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            vals = [str(c) if c is not None else "" for c in row]
            if any(v for v in vals):
                rows.append(" | ".join(vals))
        if rows:
            parts.append(f"--- {sheet_name} ---\n" + "\n".join(rows))
    return "\n\n".join(parts)


EXT_PARSERS = {
    ".pdf": parse_pdf,
    ".docx": parse_docx,
    ".doc": parse_docx,
    ".txt": parse_txt,
    ".md": parse_txt,
    ".csv": parse_txt,
    ".xlsx": parse_excel,
    ".xls": parse_excel,
}


def parse_file(file_path: str) -> dict:
    """解析单个文件，返回标准 raw_item 格式。"""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"文件不存在: {file_path}")

    ext = os.path.splitext(file_path)[1].lower()
    parser = EXT_PARSERS.get(ext)
    if parser is None:
        raise ValueError(f"不支持的文件格式: {ext}，支持: {list(EXT_PARSERS.keys())}")

    raw_text = parser(file_path)
    if not raw_text:
        logger.warning("文件内容为空: %s", file_path)

    return {
        "title": _strip_name(file_path),
        "url": os.path.abspath(file_path),
        "source": "local_file",
        "raw_text": raw_text,
        "publish_date": _file_mtime(file_path),
        "collected_at": datetime.now().isoformat(),
        "file_type": ext,
        "file_name": os.path.basename(file_path),
    }


def parse_files(file_paths: list[str]) -> list[dict]:
    """批量解析文件，返回 raw_items 列表。跳过失败的文件。"""
    items = []
    errors = []
    for path in file_paths:
        try:
            items.append(parse_file(path))
        except Exception as e:
            logger.warning("文件解析失败 [%s]: %s", path, e)
            errors.append({"file_path": path, "error": str(e)})
    if errors and not items:
        raise RuntimeError(f"所有文件解析失败: {errors}")
    return items
