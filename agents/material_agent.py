# ============================================================
# MaterialAgent — 材料辅助 Agent
# 文件：agents/material_agent.py
# 类名：MaterialAgent
# 版本：v0.1（Day 2 初版）
#
# 职责：
#   - 根据项目信息、用户信息、竞赛要求，生成申报材料内容
#   - 支持 21 种材料模板，覆盖大挑、小挑、互联网+、通用竞赛
#   - 调用 LLM 生成结构化文本（markdown 格式）
#   - 输出材料清单、填写注意事项、格式规范
#
# 依赖：
#   - config/material_prompts.yaml（Prompt 模板配置）
#   - config/config.yaml（全局配置）
# ============================================================

import yaml
import os
import re
import json
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openai import OpenAI


# ============================================================
# MaterialAgent 类
# ============================================================

class MaterialAgent:
    """
    材料辅助 Agent

    根据用户提供的项目信息、团队信息、竞赛要求，
    选择合适的 Prompt 模板，调用 LLM 生成申报材料内容。

    支持 4 大类 21 种材料模板：
      - 挑战杯系列（大挑）：6 种
      - 挑战杯系列（小挑）：4 种
      - 中国国际大学生创新大赛：5 种
      - 通用竞赛模板：6 种
    """

    # ---- 合法的 material_type 枚举 ----
    VALID_MATERIAL_TYPES = {
        # 挑战杯-大挑
        "challenge_cup_grand_application",
        "challenge_cup_grand_paper_natural",
        "challenge_cup_grand_report_social",
        "challenge_cup_grand_report_invention",
        "challenge_cup_grand_ppt",
        "challenge_cup_grand_checklist",
        # 挑战杯-小挑（创业计划竞赛）
        "challenge_cup_business_plan",
        "challenge_cup_business_ppt",
        "challenge_cup_business_application",
        "challenge_cup_business_checklist",
        # 中国国际大学生创新大赛
        "innovation_contest_business_plan",
        "innovation_contest_ppt",
        "innovation_contest_application_form",
        "innovation_contest_video_script",
        "innovation_contest_checklist",
        # 通用模板
        "generic_personal_resume",
        "generic_application_form",
        "generic_project_report",
        "generic_ppt",
        "generic_team_description",
        "generic_budget",
        "generic_schedule",
        "generic_personal_resume",
    }

    # ---- 合法的输出格式 ----
    VALID_OUTPUT_FORMATS = {"markdown", "json", "text"}

    # ---- 合法的状态值 ----
    VALID_STATUSES = {"success", "failed", "partial", "need_input", "skipped"}

    # ---- 竞赛名称 → 默认材料类型映射（推断用）----
    # key: 竞赛名称关键词（模糊匹配）
    # value: (默认 material_type, 可选 material_type 列表)
    COMPETITION_DEFAULT_MATERIAL = {
        # 大挑
        "课外学术科技": (
            "challenge_cup_grand_application",
            ["challenge_cup_grand_application", "challenge_cup_grand_paper_natural",
             "challenge_cup_grand_report_social", "challenge_cup_grand_report_invention",
             "challenge_cup_grand_ppt", "challenge_cup_grand_checklist"],
        ),
        "大挑": (
            "challenge_cup_grand_application",
            ["challenge_cup_grand_application", "challenge_cup_grand_paper_natural",
             "challenge_cup_grand_report_social", "challenge_cup_grand_report_invention",
             "challenge_cup_grand_ppt", "challenge_cup_grand_checklist"],
        ),
        # 小挑
        "创业计划竞赛": (
            "challenge_cup_business_plan",
            ["challenge_cup_business_plan", "challenge_cup_business_ppt",
             "challenge_cup_business_application", "challenge_cup_business_checklist"],
        ),
        "小挑": (
            "challenge_cup_business_plan",
            ["challenge_cup_business_plan", "challenge_cup_business_ppt",
             "challenge_cup_business_application", "challenge_cup_business_checklist"],
        ),
        # 互联网+
        "中国国际大学生创新": (
            "innovation_contest_business_plan",
            ["innovation_contest_business_plan", "innovation_contest_ppt",
             "innovation_contest_application_form", "innovation_contest_video_script",
             "innovation_contest_checklist"],
        ),
        "互联网+": (
            "innovation_contest_business_plan",
            ["innovation_contest_business_plan", "innovation_contest_ppt",
             "innovation_contest_application_form", "innovation_contest_video_script",
             "innovation_contest_checklist"],
        ),
        "互联网加": (
            "innovation_contest_business_plan",
            ["innovation_contest_business_plan", "innovation_contest_ppt",
             "innovation_contest_application_form", "innovation_contest_video_script",
             "innovation_contest_checklist"],
        ),
        # 大创
        "大创": (
            "generic_application_form",
            ["generic_application_form", "generic_project_report", "generic_ppt",
             "generic_team_description", "generic_budget", "generic_schedule"],
        ),
        "大学生创新创业训练": (
            "generic_application_form",
            ["generic_application_form", "generic_project_report", "generic_ppt",
             "generic_team_description", "generic_budget", "generic_schedule"],
        ),
        # 三创赛
        "三创": (
            "generic_project_report",
            ["generic_application_form", "generic_project_report", "generic_ppt",
             "generic_team_description", "generic_budget", "generic_schedule"],
        ),
    }

    # ============================================================
    # __init__ — 初始化 Agent
    # ============================================================

    def __init__(self, config: dict):
        """
        初始化 MaterialAgent

        Args:
            config: 配置字典，可包含以下字段：
                - model: 模型配置（provider, name, temperature, max_tokens）
                - api: API 配置（key, base_url, timeout）
                - prompt_config_path: Prompt 模板 YAML 文件路径
                - default_style: 默认语言风格
                - default_language: 默认语言
        """
        self.config = config

        # 模型配置
        llm_config = config.get("llm", {})
        model_config = config.get("model", {}) or llm_config
        self.model_provider = model_config.get("provider", "")
        self.model_name = model_config.get("name", "") or model_config.get("model", "")
        self.temperature = model_config.get("temperature", 0.7)
        self.max_tokens = model_config.get("max_tokens", 4096)

        # API 配置
        api_config = config.get("api", {})
        api_key_env = llm_config.get("api_key_env", "DEEPSEEK_API_KEY")
        self.api_key = api_config.get("key", "") or os.environ.get(api_key_env, "")
        self.api_base_url = api_config.get("base_url", "") or llm_config.get("base_url", "")
        self.api_timeout = api_config.get("timeout", llm_config.get("timeout", 60))
        self.max_retry = api_config.get(
            "max_retry", config.get("agent", {}).get("max_retry", 3)
        )

        # Agent 配置
        agent_config = config.get("agent", {}).get("material_agent", {})
        self.default_style = agent_config.get("default_style", "formal")
        self.default_language = agent_config.get("default_language", "zh-CN")

        # Prompt 模板路径
        prompt_path = agent_config.get(
            "prompt_config_path",
            "./config/material_prompts.yaml"
        )
        self.prompt_config_path = self._resolve_path(prompt_path)

        # 输出目录
        storage_config = config.get("storage", {})
        self.output_dir = self._resolve_path(
            storage_config.get("output_path", "./data/output")
        )
        self.temp_dir = self._resolve_path(
            storage_config.get("temp_path", "./data/temp")
        )

        # 确保输出目录存在
        os.makedirs(self.output_dir, exist_ok=True)
        os.makedirs(self.temp_dir, exist_ok=True)

        # 加载 Prompt 模板
        self.prompt_templates = {}
        self.global_config = {}
        self._load_prompts()

    # ============================================================
    # run — 统一外部入口
    # ============================================================

    def run(self, input_data: dict) -> dict:
        """
        MaterialAgent 统一调用入口

        Args:
            input_data: 统一输入格式的字典，结构：
                {
                  "task_id": "...",
                  "user_input": "...",
                  "task_type": "...",
                  "user_profile": {...},
                  "context": {...},
                  "input_data": {
                      "project_info": {...},
                      "user_profile": {...},
                      "competition_info": {...},
                      "material_type": "...",
                      "style": "formal"
                  },
                  "history": [...],
                  "required_output": "markdown",
                  "metadata": {...}
                }

        Returns:
            统一输出格式的字典，结构：
                {
                  "task_id": "...",
                  "agent_name": "material_agent",
                  "status": "success|failed|partial|need_input|skipped",
                  "data": {
                      "material_type": "...",
                      "material_name": "...",
                      "content": {...},
                      "format_spec": {...},
                      "checklist": [...],
                      "suggestions": [...]
                  },
                  "message": "...",
                  "error": null | {...},
                  "next_action": null | "...",
                  "metadata": {...}
                }
        """
        task_id = input_data.get("task_id", "")

        # Step 1: 校验输入
        validation_error = self.validate_input(input_data)
        if validation_error:
            return {
                "task_id": task_id,
                "agent_name": "material_agent",
                "status": "failed",
                "data": {},
                "message": "Input validation failed.",
                "error": validation_error,
                "next_action": None,
                "metadata": self._make_metadata(),
            }

        # Step 2: 处理核心业务逻辑
        try:
            result = self.process(input_data)

            # 如果 process 返回 need_input，透传
            if isinstance(result, dict) and result.get("_status") == "need_input":
                return {
                    "task_id": task_id,
                    "agent_name": "material_agent",
                    "status": "need_input",
                    "data": {},
                    "message": result.get("_message", "需要用户补充信息以确定材料类型。"),
                    "error": None,
                    "next_action": "ask_user",
                    "metadata": self._make_metadata(),
                }

            return {
                "task_id": task_id,
                "agent_name": "material_agent",
                "status": "success",
                "data": result,
                "message": f"Material '{result.get('material_type', '')}' generated successfully.",
                "error": None,
                "next_action": None,
                "metadata": self._make_metadata(),
            }

        except Exception as e:
            # 任何异常都要捕获，不允许系统崩溃
            return {
                "task_id": task_id,
                "agent_name": "material_agent",
                "status": "failed",
                "data": {},
                "message": "Material generation failed.",
                "error": {
                    "error_type": type(e).__name__,
                    "error_message": str(e),
                    "suggestion": "请检查输入数据是否完整，或尝试更换 material_type。",
                },
                "next_action": None,
                "metadata": self._make_metadata(),
            }

    # ============================================================
    # validate_input — 校验输入格式
    # ============================================================

    def validate_input(self, input_data: dict) -> dict | None:
        """
        校验输入是否合法

        检查项：
          1. input_data 子字典是否存在
          2. material_type 如果提供了，是否在 21 种合法类型中
             （如果未提供，交由 process() 推断）
          3. 必要的 project_info 是否存在

        Returns:
            None 表示校验通过，dict 表示校验失败的 error 信息
        """
        # 检查 input_data 子字典
        inner = input_data.get("input_data")
        if not inner or not isinstance(inner, dict):
            return {
                "error_type": "ValidationError",
                "error_message": "Missing or invalid 'input_data' field.",
                "suggestion": "请提供包含 project_info、material_type 等字段的 input_data。",
            }

        # 检查 material_type（如果提供了，验证合法性；未提供则留给推断）
        material_type = inner.get("material_type", "")
        if material_type and material_type not in self.VALID_MATERIAL_TYPES:
            return {
                "error_type": "ValidationError",
                "error_message": f"Invalid material_type: '{material_type}'.",
                "suggestion": f"合法的 material_type 值：{sorted(self.VALID_MATERIAL_TYPES)}",
            }

        # 检查 project_info（大部分模板需要）
        project_info = inner.get("project_info")
        if not project_info or not isinstance(project_info, dict):
            return {
                "error_type": "ValidationError",
                "error_message": "Missing or invalid 'project_info' in input_data.",
                "suggestion": "请提供项目基本信息（project_info），至少包含 project_name。",
            }

        project_name = project_info.get("project_name", "")
        if not project_name:
            return {
                "error_type": "ValidationError",
                "error_message": "Missing required field: 'project_info.project_name'.",
                "suggestion": "请提供项目名称。",
            }

        # 校验通过
        return None

    # ============================================================
    # process — 核心业务逻辑
    # ============================================================

    def process(self, input_data: dict) -> dict:
        """
        核心处理流程：
          1. 提取输入数据
          2. 推断或验证 material_type（未指定时自动匹配竞赛类型）
          3. 选择 Prompt 模板
          4. 填充占位符 → 构建完整 User Prompt
          5. 调用 LLM 生成内容
          6. 结构化解析 LLM 返回结果
          7. 组装最终输出 + 保存文件

        Args:
            input_data: 统一输入格式的字典

        Returns:
            材料生成结果 dict，包含 content、format_spec、checklist、suggestions
        """
        inner = input_data.get("input_data", {})

        material_type = inner.get("material_type", "")
        style = inner.get("style", self.default_style)
        output_format = input_data.get("required_output", "markdown")

        # ---- Step 1: 提取数据 ----
        project_info = inner.get("project_info", {})
        user_profile = inner.get("user_profile", {})
        competition_info = inner.get("competition_info", {})
        requirements = inner.get("requirements", {})

        # ---- Step 2: 推断或验证 material_type ----
        if not material_type:
            user_input = input_data.get("user_input", "")
            inference = self._infer_material_type(competition_info, user_input)
            material_type = inference["material_type"]

            if not material_type:
                # 无法推断，返回 need_input
                return {
                    "_status": "need_input",
                    "_message": (
                        f"{inference['message']}\n"
                        f"请从以下选项中选择一个 material_type：\n"
                        + "\n".join(f"  - {t}" for t in sorted(self.VALID_MATERIAL_TYPES))
                    ),
                }

            # 推断成功，记录来源
            inference_note = f"（自动推断：{inference['message']}）"
        else:
            inference_note = ""

        # ---- Step 3: 预检关键字段，缺太多则追问 ----
        missing_check = self._check_missing_fields(project_info, material_type, user_profile)
        if missing_check["need_input"]:
            return {
                "_status": "need_input",
                "_message": missing_check["message"],
            }

        # ---- Step 4: 选择 Prompt 模板 ----
        if material_type not in self.prompt_templates:
            return {
                "_status": "need_input",
                "_message": f"material_type '{material_type}' 的 Prompt 模板未加载。",
            }
        prompt_template = self.prompt_templates[material_type]
        system_prompt = self._build_system_prompt(material_type, style)
        user_prompt = self._build_user_prompt(
            prompt_template,
            project_info,
            user_profile,
            competition_info,
            requirements,
        )

        # ---- Step 5: 调用 LLM 生成 ----
        output_sections = prompt_template.get("output_sections", [])
        llm_response = self._call_llm(
            system_prompt, user_prompt,
            output_sections=output_sections,
            material_type=material_type,
        )

        # ---- Step 6: 解析 LLM 返回 ----
        parsed_content = self._parse_llm_response(
            llm_response,
            prompt_template.get("output_sections", []),
        )

        # ---- Step 7: 组装输出数据 ----
        result = {
            "material_type": material_type,
            "material_name": prompt_template.get("name", ""),
            "competition_name": (
                competition_info.get("competition_name")
                or project_info.get("project_name")
                or project_info.get("title")
                or "竞赛申报材料"
            ),
            "inference_note": inference_note,
            "content": parsed_content,
            "format_spec": prompt_template.get("format_spec", {}),
            "checklist": prompt_template.get("checklist", []),
            "suggestions": prompt_template.get("tips", []),
        }

        # ---- Step 6: 保存输出文件 ----
        task_id = input_data.get("task_id", "")
        output_format = input_data.get("required_output", "markdown")
        saved_paths = self._save_output(result, task_id, output_format)
        result["_saved_files"] = saved_paths

        return result

    # ============================================================
    # ============================================================
    # _check_missing_fields — 预检关键字段
    # ============================================================

    def _check_missing_fields(self, project_info: dict, material_type: str,
                               user_profile: dict) -> dict:
        """汇总缺失字段，全部可选，用户自主决定提供哪些。"""
        if "checklist" in material_type or material_type in (
            "generic_budget",
            "generic_team_description",
            "generic_schedule",
            "generic_personal_resume",
        ):
            return {"need_input": False, "message": ""}

        missing = []
        summary = project_info.get("summary", "")
        background = project_info.get("background", "")
        approach = project_info.get("technical_approach", "") or project_info.get("solution", "")
        innovations = project_info.get("innovation_points", [])
        team = project_info.get("team_members", [])
        advisor = project_info.get("advisor", {})

        if not summary or len(str(summary)) < 10:
            missing.append("项目简介（一两句话说一下项目是做什么的）")
        if not background or len(str(background)) < 10:
            missing.append("项目背景（要解决什么问题）")
        if not approach:
            missing.append("技术方案或解决方案（怎么实现的）")
        if not innovations or len(innovations) == 0:
            missing.append("创新点（有什么不同）")
        if not team or len(team) == 0:
            missing.append("团队成员姓名和分工")
        if not advisor:
            missing.append("指导教师姓名和研究方向")
        if not user_profile or not user_profile.get("name", ""):
            missing.append("你的姓名")

        # 检查用户是否已提供了任何信息（超过项目名就算有输入）
        has_any_info = bool(
            summary or background or approach or innovations or team or advisor
        )

        if not has_any_info and missing:
            msg = "以下信息可以帮你生成更完整的材料，全部可选，你提供多少我填多少：\n\n"
            for i, m in enumerate(missing, 1):
                msg += f"  {i}. {m}\n"
            msg += "\n回复你想提供的即可，不想提供的可以跳过，或直接说「生成」。"
            return {"need_input": True, "message": msg}

        # 用户已提供部分信息 → 直接生成，不再追问
        return {"need_input": False, "message": ""}

    # ============================================================
    # _infer_material_type — 推断材料类型
    # ============================================================

    def _infer_material_type(self, competition_info: dict, user_input: str) -> dict:
        """
        当 material_type 未明确指定时，根据竞赛名称推断默认材料类型。

        推断规则：
          1. 优先匹配 competition_info.competition_name
          2. 其次匹配 competition_info.competition_type
          3. 最后匹配 user_input 中的关键词

        Args:
            competition_info: 竞赛信息字典
            user_input: 用户原始输入文本

        Returns:
            {
              "material_type": "推断出的类型" | None,
              "confidence": "exact" | "inferred" | "none",
              "alternatives": ["备选类型1", ...],
              "message": "说明信息"
            }
        """
        # 候选匹配文本（按优先级）
        candidates = []
        # 兼容 title 和 competition_name 两种传法
        comp_name = competition_info.get("competition_name") or competition_info.get("title", "")
        comp_type = competition_info.get("competition_type", "")
        if comp_name:
            candidates.append(comp_name)
        if comp_type and comp_type != comp_name:
            candidates.append(comp_type)
        if user_input:
            candidates.append(user_input)

        # 遍历候选文本，匹配竞赛关键词
        for text in candidates:
            for keyword, (default_type, alternatives) in self.COMPETITION_DEFAULT_MATERIAL.items():
                if keyword in text:
                    # 精确匹配到具体模板名
                    # 检查 user_input 中是否有更具体的材料要求
                    specific = self._detect_specific_material(user_input, alternatives)
                    if specific:
                        return {
                            "material_type": specific,
                            "confidence": "exact",
                            "alternatives": alternatives,
                            "message": f"根据'{keyword}'匹配到竞赛类型，并根据具体需求选择材料模板。",
                        }
                    return {
                        "material_type": default_type,
                        "confidence": "inferred",
                        "alternatives": alternatives,
                        "message": f"根据'{keyword}'推断为{default_type}。如需其他材料类型，请在 material_type 中明确指定。",
                    }

        # 通用兜底：竞赛名称含"大赛"/"竞赛"/"杯"，默认用通用申报书
        for keyword in ["大赛", "竞赛", "杯"]:
            if keyword in comp_name or keyword in user_input:
                alternatives = ["generic_application_form", "generic_project_report",
                               "generic_ppt", "generic_team_description",
                               "generic_budget", "generic_schedule"]
                return {
                    "material_type": "generic_application_form",
                    "confidence": "inferred",
                    "alternatives": alternatives,
                    "message": f"未匹配到特定竞赛类型，根据'{keyword}'推断为通用申报书。如需其他材料类型，请明确指定。",
                }

        # 无法推断
        return {
            "material_type": None,
            "confidence": "none",
            "alternatives": [],
            "message": "无法从竞赛信息中推断材料类型。请明确指定 material_type。",
        }

    def _detect_specific_material(self, user_input: str, alternatives: list) -> str | None:
        """从用户输入中检测是否指定了具体材料类型"""
        material_keywords = {
            "申报书": ["application", "application_form"],
            "申报表": ["application", "application_form"],
            "报名表": ["application", "application_form"],
            "商业计划书": ["business_plan"],
            "计划书": ["business_plan", "project_report"],
            "论文": ["paper_natural"],
            "调查报告": ["report_social"],
            "研究报告": ["report_invention"],
            "科技发明": ["report_invention"],
            "PPT": ["ppt"],
            "路演": ["ppt"],
            "答辩": ["ppt"],
            "视频脚本": ["video_script"],
            "视频": ["video_script"],
            "材料清单": ["checklist"],
            "支撑材料": ["checklist"],
            "团队分工": ["team_description"],
            "分工": ["team_description"],
            "预算": ["budget"],
            "经费": ["budget"],
            "进度": ["schedule"],
            "时间规划": ["schedule"],
        }

        for word, type_hints in material_keywords.items():
            if word in user_input:
                for hint in type_hints:
                    for alt in alternatives:
                        if hint in alt:
                            return alt

        return None

    # ============================================================
    # _build_system_prompt — 构建 System Prompt
    # ============================================================

    def _build_system_prompt(self, material_type: str, style: str) -> str:
        """
        拼接全局 System Prompt + 模板专属 System Prompt

        Args:
            material_type: 材料类型标识
            style: 语言风格（formal / casual）

        Returns:
            完整的 System Prompt 字符串
        """
        # 全局通用 System Prompt
        global_system = self.global_config.get("common_system_prompt", "")
        global_system = global_system.replace("{agent_role}", self.global_config.get("agent_role", "竞赛材料辅助助手"))

        # 模板专属 System Prompt
        template = self.prompt_templates.get(material_type, {})
        template_system = template.get("system", "")

        # 风格约束（如正式风格需要额外强调）
        style_guidance = ""
        if style == "formal":
            style_guidance = (
                "\n\n## 语言风格要求\n"
                "- 使用正式、学术化的书面语言\n"
                "- 避免口语化表达和网络用语\n"
                "- 引用数据和事实时保持客观中立\n"
                "- 使用第三人称或正式的第一人称，避免随意语气"
            )

        return global_system + "\n\n" + template_system + style_guidance

    # ============================================================
    # _build_user_prompt — 构建 User Prompt
    # ============================================================

    def _build_user_prompt(
        self,
        prompt_template: dict,
        project_info: dict,
        user_profile: dict,
        competition_info: dict,
        requirements: dict,
    ) -> str:
        """
        用实际数据填充 Prompt 模板中的占位符

        占位符替换规则：
          - {project_name} → project_info.project_name
          - {background} → project_info.background
          - {team_members} → 格式化为团队成员列表文本
          - 未提供数据的占位符 → "[待补充]"

        Args:
            prompt_template: 选中的 Prompt 模板配置
            project_info: 项目信息
            user_profile: 用户信息
            competition_info: 竞赛信息

        Returns:
            填充后的 User Prompt 字符串
        """
        user_template = prompt_template.get("user_template", "")

        # 构建替换变量字典
        variables = {}

        # ---- 项目信息映射 ----
        for key in [
            "project_name", "category", "summary", "background",
            "technical_approach", "innovation_points", "current_progress",
            "achievements", "application_value", "expected_results",
            "design_purpose", "solution", "product_description",
            "ip_status", "target_market", "market_size",
            "competitor_analysis", "business_model", "revenue_model",
            "marketing_strategy", "financial_data", "funding_needed",
            "funding_usage", "social_value", "slogan", "company_info",
            "is_registered", "license_number", "project_type",
            "key_results", "experimental_plan", "test_results",
            "problems_and_improvements", "competitor_comparison",
            "application_status", "existing_materials",
            "survey_purpose", "survey_method", "survey_subjects",
            "survey_time", "survey_location", "key_findings",
            "subject_category", "author_info",
            "literature_review", "feasibility", "risk_analysis",
            "schedule", "budget", "budget_items",
            # 大挑申报书特有字段
            "patent_status", "display_forms", "tech_transfer_method",
            "development_stage", "display_notes", "applicant_degree",
            # 小挑/互联网+ 特有字段
            "operation_model", "cost_structure", "financial_summary",
            "competition_level",
            # 视频脚本特有字段
            "pain_point_scene", "product_demo", "team_description",
        ]:
            val = project_info.get(key, "")
            variables[key] = self._format_value(val)

        # ---- 用户信息映射 ----
        variables["user_name"] = user_profile.get("name", "")
        variables["user_major"] = user_profile.get("major", "")
        variables["user_grade"] = user_profile.get("grade", "")
        variables["user_college"] = user_profile.get("college", "")
        variables["user_phone"] = user_profile.get("phone", "")
        variables["user_email"] = user_profile.get("email", "")
        variables["user_skills"] = self._format_value(user_profile.get("skills", []))
        variables["user_interests"] = self._format_value(user_profile.get("interests", []))
        variables["user_experience"] = self._format_value(
            user_profile.get("competition_experience", user_profile.get("experience", []))
        )

        # ---- 竞赛信息映射 ----
        # 兼容 title 和 competition_name 两种传法
        variables["competition_name"] = competition_info.get("competition_name") or competition_info.get("title", "")
        variables["track"] = competition_info.get("track", "")
        variables["group"] = competition_info.get("group", "")
        variables["deadline"] = competition_info.get("deadline", "")
        variables["organizer"] = competition_info.get("organizer", "")
        variables["competition_type"] = competition_info.get("competition_type", "")
        variables["competition_timeline"] = competition_info.get("competition_timeline", "")
        variables["competition_level"] = competition_info.get("competition_level", variables.get("competition_level", ""))

        # ---- 特殊字段处理 ----
        variables["college"] = project_info.get("college", user_profile.get("college", ""))
        variables["field"] = project_info.get("field", project_info.get("tech_field", ""))
        variables["sub_type"] = project_info.get("sub_type", project_info.get("category", ""))
        variables["applicant_type"] = project_info.get("applicant_type", "集体")
        variables["applicant_name"] = project_info.get("applicant_name", project_info.get("project_name", ""))
        variables["submission_date"] = project_info.get("submission_date", competition_info.get("deadline", ""))
        variables["topic_group"] = project_info.get("topic_group", "")
        variables["subject_area"] = project_info.get("subject_area", "")
        variables["tech_field"] = project_info.get("tech_field", project_info.get("field", ""))

        # ---- 团队成员格式化 ----
        team_members = project_info.get("team_members", [])
        variables["team_members"] = self._format_team_members(team_members)
        variables["team_leader"] = self._format_team_leader(team_members)
        variables["team_name"] = project_info.get("team_name", variables["project_name"] + "团队")
        variables["team_size"] = str(len(team_members)) if team_members else ""
        variables["team_established_date"] = project_info.get("team_established_date", "")
        variables["team_availability"] = project_info.get("team_availability", "[待补充：团队成员每周可投入时间]")
        variables["communication_frequency"] = project_info.get("communication_frequency", "每周一次线下例会 + 线上日常沟通")
        variables["task_allocation"] = project_info.get("task_allocation", "按模块分工，使用协作工具跟踪进度")
        variables["decision_mechanism"] = project_info.get("decision_mechanism", "负责人统筹 + 集体讨论表决")

        # ---- 指导教师格式化 ----
        advisor = project_info.get("advisor", {})
        variables["advisor"] = self._format_advisor(advisor)

        # ---- 成果格式化 ----
        achievements = project_info.get("achievements", [])
        variables["achievements_detail"] = self._format_achievements(achievements)

        # ---- 经费预算格式化 ----
        budget = project_info.get("budget", {})
        variables["budget"] = self._format_budget(budget)
        variables["funding_source"] = budget.get("funding_source", project_info.get("funding_source", "[待补充]"))
        variables["project_duration"] = project_info.get("project_duration", "[待补充]")

        # ---- requirements 字段（规范 12.4 定义）----
        variables["requirements"] = self._format_value(requirements)

        # ---- 进度安排格式化 ----
        schedule = project_info.get("schedule", "")
        variables["schedule"] = self._format_value(schedule)

        # ---- 当前阶段 ----
        variables["current_phase"] = project_info.get("current_phase", "选题与团队组建")

        # 执行替换
        result = user_template
        for key, value in variables.items():
            placeholder = "{" + key + "}"
            if placeholder in result:
                result = result.replace(placeholder, str(value))

        # 处理未替换的占位符
        remaining = re.findall(r'\{(\w+)\}', result)
        for placeholder_key in remaining:
            result = result.replace(
                "{" + placeholder_key + "}",
                f"[待补充：{placeholder_key}]"
            )

        return result

    # ============================================================
    # _call_llm — 调用 LLM（占位实现，Day 3-5 填入真实 API）
    # ============================================================

    def _call_llm(self, system_prompt: str, user_prompt: str,
                   output_sections: list = None, material_type: str = "") -> str:
        """
        调用 LLM API 生成材料内容。

        如果 API 配置有效（api_key 和 base_url 均非空），调用真实 LLM；
        否则回退到 mock 模式。

        Args:
            system_prompt: System Prompt
            user_prompt: User Prompt（已填充占位符）
            output_sections: 预期的输出章节定义（mock 模式用）
            material_type: 材料类型标识（mock 模式用）

        Returns:
            LLM 生成的文本内容
        """
        # ---- 真实 API 调用 ----
        if not self.api_key or not self.api_base_url:
            has_key = "YES" if self.api_key else "NO - set DEEPSEEK_API_KEY"
            has_url = self.api_base_url or "NOT SET"
            has_model = self.model_name or "NOT SET"
            return (
                f"# [Config Error] Cannot generate material\n\n"
                f"> API Key: {has_key}\n"
                f"> Base URL: {has_url}\n"
                f"> Model: {has_model}\n\n"
                f"Please check config.yaml or DEEPSEEK_API_KEY env var."
            )
        if self.api_key and self.api_base_url:
            try:
                client = OpenAI(
                    api_key=self.api_key,
                    base_url=self.api_base_url,
                    timeout=self.api_timeout,
                )

                response = client.chat.completions.create(
                    model=self.model_name,
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt},
                    ],
                    temperature=self.temperature,
                    max_tokens=self.max_tokens,
                )

                content = response.choices[0].message.content
                usage = response.usage
                if usage:
                    print(f"    [LLM] {self.model_provider}/{self.model_name} "
                          f"tokens: {usage.total_tokens} (in:{usage.prompt_tokens} out:{usage.completion_tokens})")
                return content

            except Exception as e:
                import traceback
                has_key = "YES" if self.api_key else "NO"
                return (
                    f"# [API Error] Material generation failed\n\n"
                    f"> API Key: {has_key}\n"
                    f"> Base URL: {self.api_base_url}\n"
                    f"> Model: {self.model_name}\n"
                    f"> Error: {type(e).__name__}: {e}\n"
                    f"> Traceback: {traceback.format_exc()[-300:]}\n\n"
                    f"Please check network and API configuration."
                )

        # ---- Mock 模式（API 未配置或调用失败时）----
        if output_sections is None:
            output_sections = []

        # 使用传入的 prompt 估算 token 量（mock 模式下仅统计）
        estimated_input_tokens = len(system_prompt + user_prompt) // 2

        lines = []
        lines.append(f"# [Mock] {material_type} — 模拟生成内容")
        lines.append("")
        lines.append(f"> Mock 模式（输入约 {estimated_input_tokens} tokens），API 配置完整后替换为 AI 生成。")
        lines.append("")

        for section in output_sections:
            if isinstance(section, dict):
                title = section.get("title", "")
                section_id = section.get("section_id", "")
                slide = section.get("slide", None)

                if slide is not None:
                    lines.append(f"## Slide {slide}: {title}")
                else:
                    lines.append(f"## {title}")

                lines.append("")
                lines.append(f"[Mock 内容] 此处为「{title}」章节的模拟填充文本。")
                lines.append(f"section_id={section_id}，接入真实 LLM 后将根据项目数据生成具体的实质内容。")
                lines.append("")
            elif isinstance(section, str):
                lines.append(f"## {section}")
                lines.append("")
                lines.append(f"[Mock 内容] 此处为「{section}」的模拟内容。")
                lines.append("")

        return "\n".join(lines)

    # ============================================================
    # _parse_llm_response — 解析 LLM 返回内容
    # ============================================================

    def _parse_llm_response(
        self,
        llm_response: str,
        output_sections: list,
    ) -> dict:
        """
        将 LLM 返回的文本解析为结构化数据。

        解析策略（按优先级）：
          1. 尝试提取 ```json ... ``` 代码块，JSON 解析
          2. 尝试将全文作为 JSON 解析
          3. 按 Markdown ## 标题分割（fallback）
          4. 以上都失败，返回全文作为单个 section

        Args:
            llm_response: LLM 返回的文本
            output_sections: 预期的输出章节定义

        Returns:
            dict: {"cover": {}, "sections": [...], "raw_text": "..."}
        """
        # ---- 策略 1: 提取 JSON 代码块 ----
        json_match = re.search(r'```(?:json)?\s*\n?(.*?)\n?\s*```', llm_response, re.DOTALL)
        if json_match:
            try:
                data = json.loads(json_match.group(1).strip())
                if isinstance(data, (dict, list)):
                    return self._parse_from_json(data, output_sections, llm_response)
            except json.JSONDecodeError:
                pass

        # ---- 策略 2: 全文 JSON 解析 ----
        stripped = llm_response.strip()
        if stripped.startswith("{") and stripped.endswith("}"):
            try:
                data = json.loads(stripped)
                return self._parse_from_json(data, output_sections, llm_response)
            except json.JSONDecodeError:
                pass

        # ---- 策略 3: Markdown ## 标题分割 ----
        sections = self._parse_from_markdown(llm_response, output_sections)
        if sections:
            return {
                "cover": {},
                "sections": sections,
                "raw_text": llm_response,
            }

        # ---- 策略 4: 全文兜底 ----
        return {
            "cover": {},
            "sections": [{
                "section_id": "full_text",
                "title": "生成内容",
                "content": llm_response.strip(),
            }],
            "raw_text": llm_response,
        }

    def _parse_from_json(self, data: dict, output_sections: list, raw_text: str) -> dict:
        """从 JSON 数据构建 sections"""
        sections = []

        # JSON 可能是 {"sections": [...]} 或直接是列表
        json_sections = data.get("sections", data if isinstance(data, list) else [])
        if isinstance(json_sections, dict):
            json_sections = [json_sections]

        for item in json_sections:
            if isinstance(item, dict):
                sections.append({
                    "section_id": item.get("section_id", item.get("title", "")),
                    "title": item.get("title", ""),
                    "content": item.get("content", ""),
                })
            elif isinstance(item, str):
                sections.append({
                    "section_id": "",
                    "title": "",
                    "content": item,
                })

        return {
            "cover": data.get("cover", {}),
            "sections": sections,
            "raw_text": raw_text,
        }

    def _parse_from_markdown(self, text: str, output_sections: list) -> list:
        """按 Markdown ## 标题分割成 sections"""
        sections = []
        current_section = None
        current_content = []

        for line in text.split("\n"):
            if line.startswith("## "):
                if current_section is not None:
                    current_section["content"] = "\n".join(current_content).strip()
                    sections.append(current_section)

                section_title = line[3:].strip()
                current_section = {
                    "section_id": self._title_to_section_id(section_title, output_sections),
                    "title": section_title,
                    "content": "",
                }
                current_content = []
            else:
                if current_section is not None:
                    current_content.append(line)

        if current_section is not None:
            current_section["content"] = "\n".join(current_content).strip()
            sections.append(current_section)

        return sections

    # ============================================================
    # _save_output — 保存输出文件
    # ============================================================

    def _save_output(self, result: dict, task_id: str, output_format: str) -> list:
        """
        将生成的材料保存到 data/output/ 目录

        文件命名规则：
          {竞赛名称}_{材料名称}.docx

        Args:
            result: process() 组装的结果 dict
            task_id: 任务编号
            output_format: 保留兼容旧调用；下载文件统一生成为 Word

        Returns:
            list: 已保存的文件路径列表
        """
        competition_name = self._safe_filename(
            str(result.get("competition_name") or "竞赛申报材料")
        )
        material_name = self._safe_filename(
            str(result.get("material_name") or "申报材料")
        )
        filename = f"{competition_name}_{material_name}.docx"
        docx_path = os.path.join(self.output_dir, filename)
        self._build_docx(result, docx_path)
        return [docx_path]

    def _safe_filename(self, value: str) -> str:
        """Return a readable Windows-safe filename component."""
        value = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", value).strip(" ._")
        value = re.sub(r"\s+", "", value)
        return value[:80] or "竞赛申报材料"

    def _build_docx(self, result: dict, output_path: str) -> None:
        """Create an editable Word document using a compact reference layout."""
        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)

        styles = document.styles
        normal = styles["Normal"]
        normal.font.name = "Microsoft YaHei"
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        normal.font.size = Pt(11)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.25

        heading_tokens = {
            "Title": (24, "0B2545", 0, 8),
            "Heading 1": (16, "2E74B5", 18, 10),
            "Heading 2": (13, "2E74B5", 14, 7),
        }
        for style_name, (size, color, before, after) in heading_tokens.items():
            style = styles[style_name]
            style.font.name = "Microsoft YaHei"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            style.font.size = Pt(size)
            style.font.color.rgb = RGBColor.from_string(color)
            style.paragraph_format.space_before = Pt(before)
            style.paragraph_format.space_after = Pt(after)

        for list_style_name in ("List Bullet", "List Number"):
            list_style = styles[list_style_name]
            list_style.font.name = "Microsoft YaHei"
            list_style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            list_style.font.size = Pt(11)
            list_style.paragraph_format.left_indent = Inches(0.375)
            list_style.paragraph_format.first_line_indent = Inches(-0.188)
            list_style.paragraph_format.space_after = Pt(4)
            list_style.paragraph_format.line_spacing = 1.25

        title = document.add_paragraph(style="Title")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.add_run(str(result.get("competition_name") or "竞赛申报材料"))

        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.paragraph_format.space_after = Pt(18)
        run = subtitle.add_run(str(result.get("material_name") or "申报材料"))
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(80, 80, 80)

        metadata = document.add_paragraph()
        metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
        metadata.paragraph_format.space_after = Pt(18)
        metadata_run = metadata.add_run(
            f"材料类型：{result.get('material_type', '')}  |  生成时间：{datetime.now().strftime('%Y-%m-%d')}"
        )
        metadata_run.font.size = Pt(9.5)
        metadata_run.font.color.rgb = RGBColor(100, 100, 100)

        content = result.get("content", {})
        sections = content.get("sections", []) if isinstance(content, dict) else []
        for item in sections:
            if not isinstance(item, dict):
                continue
            heading = str(item.get("title") or "").strip()
            body = str(item.get("content") or "").strip()
            if heading:
                document.add_heading(heading, level=1)
            for block in re.split(r"\n\s*\n", body):
                block = block.strip()
                if block:
                    document.add_paragraph(block)

        checklist = result.get("checklist", [])
        if checklist:
            document.add_heading("准备清单", level=1)
            for item in checklist:
                document.add_paragraph(str(item), style="List Bullet")

        suggestions = result.get("suggestions", [])
        if suggestions:
            document.add_heading("填写建议", level=1)
            for tip in suggestions:
                document.add_paragraph(str(tip), style="List Bullet")

        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.add_run("AI 生成初稿，请根据竞赛官方要求人工核对并完善")
        footer_run.font.size = Pt(8.5)
        footer_run.font.color.rgb = RGBColor(120, 120, 120)

        document.save(output_path)

    # ============================================================
    # 辅助方法
    # ============================================================

    def _load_prompts(self):
        """加载 Prompt 模板 YAML 文件"""
        try:
            with open(self.prompt_config_path, "r", encoding="utf-8") as f:
                data = yaml.safe_load(f)

            self.global_config = data.get("global", {})

            for key, value in data.items():
                if key not in ("global", "material_type_index"):
                    self.prompt_templates[key] = value

        except FileNotFoundError:
            raise FileNotFoundError(
                f"Prompt template file not found: {self.prompt_config_path}. "
                "请确保 config/material_prompts.yaml 文件存在。"
            )
        except yaml.YAMLError as e:
            raise ValueError(f"Failed to parse prompt template YAML: {e}")

    def _resolve_path(self, path: str) -> str:
        """解析相对路径为绝对路径"""
        if os.path.isabs(path):
            return path
        # 相对于项目根目录
        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        return os.path.normpath(os.path.join(base_dir, path))

    def _make_metadata(self) -> dict:
        """生成 metadata 信息"""
        return {
            "agent_version": "0.1",
            "executed_at": datetime.now().isoformat(),
            "model_provider": self.model_provider,
            "model_name": self.model_name,
        }

    def _format_value(self, value) -> str:
        """将 Python 对象格式化为 prompt 可用的字符串"""
        if value is None:
            return "[待补充]"
        if isinstance(value, list):
            return "\n".join([f"- {item}" for item in value])
        if isinstance(value, dict):
            return "\n".join([f"- {k}: {v}" for k, v in value.items()])
        if isinstance(value, str) and value.strip() == "":
            return "[待补充]"
        return str(value)

    def _format_team_members(self, team_members: list) -> str:
        """将团队成员列表格式化为可读文本"""
        if not team_members:
            return "[待补充：团队成员信息]"

        lines = []
        for i, member in enumerate(team_members, 1):
            if isinstance(member, dict):
                name = member.get("name", "")
                role = member.get("role", "")
                major = member.get("major", "")
                grade = member.get("grade", "")
                contributions = member.get("contributions", "")
                lines.append(
                    f"{i}. {name} — {role}"
                    + (f"（{major} {grade}）" if major or grade else "")
                    + (f"：{contributions}" if contributions else "")
                )
            else:
                lines.append(f"{i}. {member}")
        return "\n".join(lines)

    def _format_team_leader(self, team_members: list) -> str:
        """提取团队负责人信息"""
        if not team_members:
            return "[待补充]"
        for member in team_members:
            if isinstance(member, dict):
                role = member.get("role", "")
                if "负责" in role or "负责人" in role or "队长" in role:
                    name = member.get("name", "")
                    major = member.get("major", "")
                    grade = member.get("grade", "")
                    return f"{name}（{major} {grade}）"
        # 默认取第一个成员
        first = team_members[0]
        if isinstance(first, dict):
            return f"{first.get('name', '')}（{first.get('major', '')} {first.get('grade', '')}）"
        return str(first)

    def _format_advisor(self, advisor) -> str:
        """格式化指导教师信息"""
        if not advisor:
            return "[待补充：指导教师信息]"
        if isinstance(advisor, dict):
            name = advisor.get("name", "")
            title = advisor.get("title", "")
            direction = advisor.get("direction", "")
            return f"{name} {title}" + (f"（研究方向：{direction}）" if direction else "")
        return str(advisor)

    def _format_achievements(self, achievements: list) -> str:
        """格式化已有成果列表"""
        if not achievements:
            return "[待补充：已有成果信息]"
        lines = []
        for item in achievements:
            if isinstance(item, dict):
                atype = item.get("type", "")
                title = item.get("title", "")
                date = item.get("date", "")
                lines.append(f"- [{atype}] {title}" + (f"（{date}）" if date else ""))
            else:
                lines.append(f"- {item}")
        return "\n".join(lines)

    def _format_budget(self, budget) -> str:
        """格式化经费预算"""
        if not budget:
            return "[待补充：经费预算信息]"
        if isinstance(budget, dict):
            total = budget.get("total", 0)
            items = budget.get("items", [])
            lines = [f"总预算：{total}元"]
            for item in items:
                if isinstance(item, dict):
                    lines.append(
                        f"- {item.get('name', '')}：{item.get('amount', 0)}元"
                    )
            return "\n".join(lines)
        return str(budget)

    def _title_to_section_id(self, title: str, output_sections: list) -> str:
        """根据标题文本匹配 section_id"""
        # 简单匹配
        for section in output_sections:
            if isinstance(section, dict):
                sec_title = section.get("title", "")
                if sec_title and (sec_title in title or title in sec_title):
                    return section.get("section_id", "")
        # fallback: 用标题生成 id
        return title.replace(" ", "_").replace("（", "").replace("）", "").lower()


# ============================================================
# 模块自测（Day 2 初版）
# ============================================================

if __name__ == "__main__":
    # ---- 辅助函数 ----
    def make_test(agent, task_id, user_input, material_type,
                  project_info, user_profile=None, competition_info=None):
        """构造测试输入并执行"""
        inner = {
            "project_info": project_info,
            "material_type": material_type,
            "style": "formal",
        }
        if user_profile:
            inner["user_profile"] = user_profile
        if competition_info:
            inner["competition_info"] = competition_info

        return agent.run({
            "task_id": task_id,
            "user_input": user_input,
            "task_type": "material_generation",
            "user_profile": user_profile or {},
            "context": {},
            "input_data": inner,
            "history": [],
            "required_output": "markdown",
            "metadata": {},
        })

    def check_result(result, label, expected_mt=None, min_sections=1):
        """检查结果并打印"""
        ok = True
        status = result.get("status", "")
        data = result.get("data", {})
        mt = data.get("material_type", "")
        sections = data.get("content", {}).get("sections", [])
        checklist = data.get("checklist", [])
        suggestions = data.get("suggestions", [])

        if status != "success" and expected_mt is not None:
            print(f"  {label}: FAIL (status={status}, expected=success)")
            return False
        if expected_mt and mt != expected_mt:
            print(f"  {label}: FAIL (material_type={mt}, expected={expected_mt})")
            return False
        if len(sections) < min_sections:
            print(f"  {label}: FAIL (sections={len(sections)}, expected>={min_sections})")
            return False

        print(f"  {label}: PASS  sections={len(sections)}  checklist={len(checklist)}  suggestions={len(suggestions)}  mt={mt}")
        return True

    # ================================================================
    print("=" * 60)
    print("MaterialAgent Day 4 自测")
    print("=" * 60)

    # 优先读 config.yaml，读不到则用空配置（mock 模式）
    try:
        with open("config/config.yaml", "r", encoding="utf-8") as f:
            test_config = yaml.safe_load(f)
        print(f"[Config] Loaded from config/config.yaml")
    except Exception:
        test_config = {
            "model": {"provider": "", "name": "", "temperature": 0.7, "max_tokens": 4096},
            "api": {"key": "", "base_url": "", "timeout": 60},
            "agent": {"material_agent": {"default_style": "formal", "prompt_config_path": "./config/material_prompts.yaml"}},
            "storage": {"output_path": "./data/output", "temp_path": "./data/temp"},
        }
        print(f"[Config] config.yaml not found, using mock mode")

    agent = MaterialAgent(test_config)
    print(f"\n[Init] Loaded {len(agent.prompt_templates)} prompt templates")

    # ---- 共享的丰富项目数据（科技发明类）----
    rich_tech_project = {
        "project_name": "智瞳——基于多光谱视觉的农作物病虫害智能监测系统",
        "category": "科技发明制作A类",
        "subject_category": "信息技术",
        "college": "计算机与人工智能学院",
        "summary": "融合多光谱成像与轻量级深度学习，实现田间作物病虫害的实时、无损、精准识别与预警",
        "background": "我国每年因病虫害造成的粮食损失超千亿元。传统人工巡检效率低、主观性强，农药滥用现象严重。现有无人机方案成本高、依赖网络，小农户难以承受。",
        "technical_approach": "多光谱相机（RGB+NIR）边缘采集→轻量级MobileNet-V3改进模型端侧推理→LoRa低功耗广域网回传→云端大数据分析平台→微信小程序推送预警",
        "innovation_points": [
            "首创可见光+近红外双模态融合的端侧病虫害识别算法，识别率92.7%",
            "自研模型压缩技术，参数量仅2.1M，可在STM32H7上实时运行",
            "离线优先架构：无网络环境下仍可独立工作8小时",
            "与浙江省农科院共建5万+标注样本的病虫害多光谱数据集"
        ],
        "current_progress": "已完成样机研发，在温州瓯海区3个农业合作社试点3个月，累计监测面积1200亩",
        "application_value": "可降低农药使用量30%以上，每亩节省成本约80元，全国推广后年经济效益超百亿",
        "patent_status": "已获专利权批准",
        "display_forms": "实物、现场演示、图片、录像",
        "tech_transfer_method": "技术许可+联合运营",
        "development_stage": "B中试阶段",
        "applicant_type": "集体",
        "applicant_name": "智瞳科技团队",
        "applicant_degree": "本科",
        "team_members": [
            {"name": "陈宇", "role": "负责人/算法研发", "major": "计算机科学与技术", "grade": "大三", "contributions": "整体架构设计、深度学习模型开发"},
            {"name": "林晓", "role": "硬件开发", "major": "电子信息工程", "grade": "大三", "contributions": "多光谱采集模块、LoRa通信模块设计"},
            {"name": "赵敏", "role": "农学顾问", "major": "农学", "grade": "研二", "contributions": "病虫害知识库、田间试验设计"},
            {"name": "张昊", "role": "软件开发", "major": "软件工程", "grade": "大三", "contributions": "云端平台、微信小程序开发"},
            {"name": "周婷", "role": "市场运营", "major": "市场营销", "grade": "大二", "contributions": "用户调研、合作社对接、商业推广"},
        ],
        "advisor": {"name": "刘建国", "title": "教授", "direction": "农业信息化与智能装备"},
        "achievements": [
            {"type": "发明专利", "title": "一种基于多光谱的农作物病虫害识别方法", "date": "2025-12"},
            {"type": "软件著作权", "title": "智瞳病虫害监测预警平台V1.0", "date": "2025-11"},
            {"type": "论文", "title": "Lightweight Multi-Spectral Fusion Network for Crop Pest Detection", "date": "2025-10"},
            {"type": "获奖", "title": "浙江省大学生科技创新大赛一等奖", "date": "2025-09"},
        ],
        "budget": {"total": 50000, "items": [
            {"name": "多光谱相机模组", "amount": 15000},
            {"name": "开发板及传感器", "amount": 10000},
            {"name": "田间试验差旅费", "amount": 12000},
            {"name": "专利申请费", "amount": 8000},
            {"name": "打印及文献费", "amount": 5000},
        ]},
        "funding_source": "导师课题经费+学校大创资助",
        "schedule": "2025.03-2025.06 方案设计与样机开发\n2025.07-2025.12 田间试验与模型优化\n2026.01-2026.06 产品迭代与推广",
        "current_phase": "产品迭代与推广",
    }

    # ---- 共享的丰富项目数据（创业类）----
    rich_biz_project = {
        "project_name": "书巢——校园二手书智能流转平台",
        "slogan": "让每一本教材找到下一位主人",
        "category": "创业实践",
        "field": "文化创意和区域合作",
        "summary": "基于智能定价算法和校内物流网络的C2C二手教材交易平台",
        "background": "高校教材年均花费超2000元/生，二手教材流转率不足15%。传统二手群交易效率低、定价随意，每年数百万本教材被当废纸处理。",
        "solution": "微信小程序+智能定价模型+校内分布式仓储+学生兼职配送",
        "product_description": "卖家扫码ISBN自动填写书信息→系统基于使用痕迹和市场需求自动定价→买家搜索/订阅目标教材→下单后校内24小时送达",
        "technical_approach": "基于历史交易数据的XGBoost动态定价模型+图像识别自动评估书况+LBS匹配最近仓储点+路径优化配送算法",
        "innovation_points": [
            "全国首创教材自动定价模型，基于20万+真实交易数据训练",
            "校内分布式仓储：利用宿舍楼闲置空间，单点覆盖半径500米",
            "绿色积分体系：每交易一本折合碳减排1.2kg，可兑换校园消费券"
        ],
        "ip_status": "已申请软件著作权1项，正在申请发明专利1项",
        "target_market": "全国高校在校生（3800万+），首期聚焦浙江省内10所高校",
        "market_size": "高校教材市场年规模约600亿元，二手流转市场约80亿元，年增长率12%",
        "competitor_analysis": "现有方案：①校园二手群（效率低、无保障）②闲鱼（运费高、周期长）③多抓鱼（不收教材）。本方案填补教材垂直领域的市场空白。",
        "business_model": "交易佣金10%+广告推广+毕业季回收增值服务",
        "revenue_model": "每笔交易抽取10%服务费（均价25元/本，佣金2.5元），预计首年交易量5万笔，营收12.5万元",
        "operation_model": "自营+众包：核心团队管理平台和算法，仓储和配送由校园合伙人负责",
        "marketing_strategy": "开学季免费收书引流→班级推广大使裂变→绿色校园公益背书→高校官方合作",
        "financial_data": "试运营2个月：注册用户2300人，交易笔数860单，交易额2.1万元",
        "cost_structure": "固定成本（服务器+办公）3000元/月；变动成本（配送补贴）1.5元/单",
        "funding_needed": "种子轮融资20万元，出让10%股权",
        "funding_usage": "技术开发40%+市场推广30%+运营储备30%",
        "social_value": "预计每年为试点高校学生节省教材费用200万元+每年减少纸张浪费约50吨+带动50+勤工助学岗位",
        "current_progress": "已在温州大学试运营2个月，正筹备浙江工业大学和杭州电子科技大学推广",
        "team_members": [
            {"name": "李悦", "role": "创始人/产品", "major": "工商管理", "grade": "大三", "contributions": "整体运营、商业模式设计"},
            {"name": "陈浩", "role": "技术负责人", "major": "软件工程", "grade": "研一", "contributions": "小程序开发、定价算法"},
            {"name": "王雨", "role": "市场推广", "major": "市场营销", "grade": "大二", "contributions": "校园推广、高校合作对接"},
            {"name": "刘洋", "role": "供应链管理", "major": "物流管理", "grade": "大三", "contributions": "仓储网络、配送体系设计"},
        ],
        "advisor": {"name": "张明华", "title": "副教授", "direction": "创业管理与电子商务"},
        "achievements": [
            {"type": "软件著作权", "title": "书巢二手书交易平台V1.0", "date": "2025-12"},
            {"type": "获奖", "title": "校创新创业大赛金奖", "date": "2025-11"},
            {"type": "合作", "title": "与温州大学后勤集团达成官方合作", "date": "2025-10"},
        ],
    }

    # ================================================================
    # 深度测试 1: 大挑申报书（表格型，字段最多）
    # ================================================================
    print(f"\n[深度测试1] challenge_cup_grand_application（大挑申报书-表格型）")
    r1 = make_test(agent, "deep_001", "帮我填写大挑科技发明类申报书",
                   "challenge_cup_grand_application", rich_tech_project)
    check_result(r1, "大挑申报书", "challenge_cup_grand_application", min_sections=5)

    # 检查是否补齐了全部 6 个表
    section_ids = [s["section_id"] for s in r1.get("data", {}).get("content", {}).get("sections", [])]
    expected_tables = ["cover", "A", "B", "C", "D", "E"]
    missing_tables = [t for t in expected_tables if t not in section_ids]
    if missing_tables:
        print(f"        缺失表格: {missing_tables}")
    else:
        print(f"        六表完整: {expected_tables}")

    # 检查占位符残留
    raw_user_prompt = ""  # 无法直接获取，通过检查最终输出
    raw_text = r1.get("data", {}).get("content", {}).get("raw_text", "")
    unstub = raw_text.count("[待补充")
    print(f"        未填充占位符: {unstub}")

    # ================================================================
    # 深度测试 2: 小挑商业计划书（长文型，10章结构）
    # ================================================================
    print(f"\n[深度测试2] challenge_cup_business_plan（小挑商业计划书-长文型）")
    r2 = make_test(agent, "deep_002", "帮我写小挑商业计划书",
                   "challenge_cup_business_plan", rich_biz_project)
    check_result(r2, "小挑商业计划书", "challenge_cup_business_plan", min_sections=8)

    section_ids_2 = [s["section_id"] for s in r2.get("data", {}).get("content", {}).get("sections", [])]
    # 10 章结构
    expected_chapters = ["executive_summary", "background", "project_intro",
                         "market", "business_model", "team", "finance",
                         "social_value", "risk", "future"]
    missing_ch = [c for c in expected_chapters if c not in section_ids_2]
    if missing_ch:
        print(f"        缺失章节: {missing_ch}")
    else:
        print(f"        十章完整: {expected_chapters}")

    # 检查执行摘要字数约束
    exec_section = next((s for s in r2.get("data", {}).get("content", {}).get("sections", [])
                        if "executive" in s.get("section_id", "")), None)
    if exec_section:
        word_count = len(exec_section.get("content", ""))
        print(f"        执行摘要字数: {word_count} (规范要求≤800字)")

    # ================================================================
    # 深度测试 3: 通用PPT（幻灯片型，slide 结构）
    # ================================================================
    print(f"\n[深度测试3] generic_ppt（通用路演PPT-幻灯片型）")
    r3 = make_test(agent, "deep_003", "帮我做项目路演PPT",
                   "generic_ppt", rich_biz_project,
                   competition_info={"competition_type": "创业型"})
    check_result(r3, "通用路演PPT", "generic_ppt", min_sections=8)

    # 检查 slide 结构
    slide_sections = [s for s in r3.get("data", {}).get("content", {}).get("sections", [])
                     if "slide" in s.get("section_id", "").lower() or "Slide" in s.get("title", "")]
    print(f"        Slide 页数: {len(slide_sections)} (规范要求≤20页)")

    # ================================================================
    # 输入校验
    # ================================================================
    print(f"\n[输入校验]")
    checks = [
        ("允许缺失 material_type",
         agent.validate_input({"input_data": {"project_info": {"project_name": "test"}}}) is None),
        ("非法 material_type",
         agent.validate_input({"input_data": {"material_type": "invalid", "project_info": {"project_name": "t"}}}) is not None),
        ("空 project_info",
         agent.validate_input({"input_data": {"material_type": "generic_ppt", "project_info": {}}}) is not None),
        ("缺失 project_name",
         agent.validate_input({"input_data": {"material_type": "generic_ppt", "project_info": {"bg": "t"}}}) is not None),
    ]
    for label, passed in checks:
        print(f"  {'PASS' if passed else 'FAIL'}  {label}")

    # ================================================================
    # 推断测试
    # ================================================================
    print(f"\n[推断测试]")
    inf_tests = [
        ("小挑", make_test(agent, "inf_01", "帮我准备小挑的商业计划书", "",
                           {"project_name": "测试项目"},
                           competition_info={"competition_name": "挑战杯中国大学生创业计划竞赛", "competition_type": "小挑"}),
         "challenge_cup_business_plan"),
        ("互联网+", make_test(agent, "inf_02", "互联网+大赛需要准备什么", "",
                              {"project_name": "测试项目"},
                              competition_info={"competition_name": "中国国际大学生创新大赛", "competition_type": "互联网+"}),
         "innovation_contest_business_plan"),
        ("大挑", make_test(agent, "inf_03", "大挑申报", "",
                           {"project_name": "测试项目"},
                           competition_info={"competition_name": "挑战杯全国大学生课外学术科技作品竞赛", "competition_type": "大挑"}),
         "challenge_cup_grand_application"),
        ("无法推断", make_test(agent, "inf_04", "帮我准备比赛材料", "",
                               {"project_name": "测试项目"}),
         None),  # 不检查 mt, 检查 need_input
    ]
    for label, r, exp_mt in inf_tests:
        if exp_mt:
            mt = r.get("data", {}).get("material_type", "")
            print(f"  {'PASS' if mt == exp_mt else 'FAIL'}  {label}: {mt}")
        else:
            print(f"  {'PASS' if r['status'] == 'need_input' else 'FAIL'}  {label}: status={r['status']}")

    # ================================================================
    # 冒烟测试：其余 18 种模板加载+process不抛异常
    # ================================================================
    print(f"\n[冒烟测试] 18 种模板 process() 不抛异常")
    smoke_pass = 0
    smoke_fail = []
    all_types = sorted(agent.VALID_MATERIAL_TYPES)
    deep_tested = {"challenge_cup_grand_application", "challenge_cup_business_plan", "generic_ppt"}

    for mt in all_types:
        if mt in deep_tested:
            continue
        try:
            r = make_test(agent, f"smoke_{mt}", f"测试{mt}", mt,
                          {"project_name": "冒烟测试项目", "background": "测试", "summary": "测试"})
            if r.get("status") == "success":
                smoke_pass += 1
            else:
                smoke_fail.append(f"{mt}(status={r.get('status')})")
        except Exception as e:
            smoke_fail.append(f"{mt}(exception={e})")

    if smoke_fail:
        for f in smoke_fail:
            print(f"  FAIL  {f}")
    else:
        print(f"  ALL {smoke_pass} PASS")

    # ================================================================
    print(f"\n{'=' * 60}")
    print("Day 4 自测完成")
    print(f"{'=' * 60}")
